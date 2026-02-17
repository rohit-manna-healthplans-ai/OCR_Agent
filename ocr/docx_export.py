from __future__ import annotations

from typing import Any, Dict
from docx import Document


def build_docx_from_result(result: Dict[str, Any], out_path: str, title: str = "OCR Output") -> None:
    doc = Document()
    doc.add_heading(title, level=1)

    meta = result.get("meta") or {}
    doc.add_paragraph(
        f"Engine: {meta.get('engine', '')} | Lang: {meta.get('lang','')} | Preset: {meta.get('preset_used', meta.get('preset_requested',''))}"
    )

    fields = result.get("fields") or {}
    if fields:
        doc.add_heading("Extracted Fields", level=2)
        t = doc.add_table(rows=1, cols=2)
        hdr = t.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"
        for k, v in fields.items():
            row = t.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)

    doc.add_heading("Extracted Text", level=2)
    text = result.get("text") or ""
    for para in text.splitlines():
        doc.add_paragraph(para)

    pages = result.get("pages") or []
    any_tables = False
    for p in pages:
        for i, tb in enumerate(p.get("tables") or []):
            any_tables = True
            doc.add_heading(f"Detected Table (page {p.get('page_index', 0)} - #{i+1})", level=2)
            cells = tb.get("cells") or []
            if not cells:
                continue
            rows = len(cells)
            cols = len(cells[0]) if rows else 0
            if rows == 0 or cols == 0:
                continue
            dt = doc.add_table(rows=rows, cols=cols)
            for r in range(rows):
                for c in range(cols):
                    dt.cell(r, c).text = str(cells[r][c] or "")

    if not any_tables:
        doc.add_heading("Detected Tables", level=2)
        doc.add_paragraph("No tables detected (best-effort).")

    doc.save(out_path)
